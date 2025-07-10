import os
import tempfile
import mimetypes
import re
import io
import chardet
import traceback
from django.conf import settings
import sys

# 设置Tesseract OCR路径（如果存在）
try:
    import pytesseract
    if os.name == 'nt':  # Windows 系统
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except ImportError:
    pytesseract = None
    print("未安装pytesseract，OCR功能将不可用")

def detect_encoding(file_path):
    """检测文件编码"""
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read(min(f.seek(0, 2), 1024 * 1024))
            f.seek(0)
            result = chardet.detect(raw_data)
            encoding = result['encoding'] if result['confidence'] > 0.7 else 'utf-8'
            return encoding
    except Exception:
        return 'utf-8'  # 默认返回UTF-8

def process_file_content(file_path, filename, content_type=None):
    """
    处理不同类型的文件内容
    
    参数:
        file_path: 文件路径
        filename: 文件名
        content_type: 文件类型
        
    返回:
        (文件内容字符串, 内容类型)
    """
    try:
        # 获取文件扩展名
        _, ext = os.path.splitext(filename.lower())
        
        # 如果未提供content_type，尝试从文件扩展名推断
        if not content_type:
            content_type = mimetypes.guess_type(filename)[0]
        
        # 打印调试信息
        # print(f"开始处理文件: {filename} (类型: {content_type}, 扩展名: {ext})")
        
        # 检查文件是否存在且可读
        if not os.path.exists(file_path) or not os.access(file_path, os.R_OK):
            return f"文件不存在或不可读: {file_path}", "text/plain"
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        # print(f"文件大小: {file_size} 字节")
        if file_size == 0:
            return "文件为空，无内容可提取", "text/plain"
        
        # 图像文件处理
        if content_type and content_type.startswith('image/'):
            try:
                from PIL import Image, ImageEnhance
                
                # 打开图像
                image = Image.open(file_path)
                width, height = image.size
                # print(f"图像尺寸: {width}x{height}, 模式: {image.mode}")
                
                # 图像预处理 - 增强对比度，使OCR更准确
                enhancer = ImageEnhance.Contrast(image)
                enhanced_image = enhancer.enhance(1.5)  # 增强对比度
                
                # 使用pytesseract执行OCR
                if pytesseract:
                    try:
                        # 先尝试中英文混合识别
                        text = pytesseract.image_to_string(enhanced_image, lang='chi_sim+eng', config='--psm 3')
                        
                        if not text.strip():
                            # 如果识别为空，尝试仅英文识别
                            print("中英文识别为空，尝试仅英文识别")
                            text = pytesseract.image_to_string(enhanced_image, lang='eng', config='--psm 3')
                            
                        if not text.strip():
                            # 如果仍为空，尝试降低DPI进行识别
                            print("标准识别为空，尝试调整DPI识别")
                            text = pytesseract.image_to_string(enhanced_image, lang='chi_sim+eng', config='--psm 1 --dpi 300')
                        
                        # print(f"OCR识别结果长度: {len(text)}")
                    except Exception as ocr_err:
                        print(f"OCR识别出错: {str(ocr_err)}")
                        traceback.print_exc()
                        text = f"OCR识别失败: {str(ocr_err)}"
                else:
                    text = "系统未安装OCR组件，无法识别图像中的文字。"
                
                # 清理文本
                text = re.sub(r'\s+', ' ', text).strip()
                
                if not text:
                    print("OCR未能识别出任何文本")
                    return "图像OCR识别未提取到文本内容", "text/plain"
                
                result = f"[图像文件内容 - OCR识别结果]\n{text}"
                # print(f"成功提取图像文本，长度: {len(result)}")
                return result, "extracted_text"
                
            except Exception as e:
                print(f"图像处理失败: {str(e)}")
                traceback.print_exc()
                return f"无法处理图像: {str(e)}", "text/plain"
                
        # PDF文件处理
        elif ext == '.pdf':
            try:
                # 尝试方法1: 使用PyMuPDF/fitz
                try:
                    import fitz
                    print("使用PyMuPDF处理PDF")
                    
                    text_content = []
                    
                    # 打开PDF文件
                    doc = fitz.open(file_path)
                    print(f"PDF页数: {doc.page_count}")
                    
                    # 遍历每一页
                    for page_num in range(doc.page_count):
                        page = doc[page_num]
                        # 提取文本
                        text = page.get_text()
                        if text.strip():
                            text_content.append(f"=== 第 {page_num + 1} 页 ===\n{text}")
                            print(f"页面 {page_num + 1} 文本长度: {len(text)}")
                    
                    # 关闭文档
                    doc.close()
                    
                    # 如果未能提取到文本，尝试OCR
                    if not text_content and pytesseract:
                        print("PyMuPDF未能提取到文本，尝试OCR")
                        from pdf2image import convert_from_path
                        from PIL import Image
                        import tempfile
                        
                        ocr_text = []
                        with tempfile.TemporaryDirectory() as path:
                            # 将PDF转换为图像
                            pdf_pages = convert_from_path(file_path, 300)
                            
                            # 对每个页面进行OCR
                            for page_num, page in enumerate(pdf_pages):
                                page_text = pytesseract.image_to_string(page, lang='chi_sim+eng')
                                if page_text.strip():
                                    ocr_text.append(f"=== 第 {page_num + 1} 页 (OCR) ===\n{page_text}")
                        
                        if ocr_text:
                            text_content = ocr_text
                            print(f"OCR成功提取PDF文本，页数: {len(ocr_text)}")
                    
                    if not text_content:
                        return "PDF文件未包含可提取的文本内容", "text/plain"
                    
                    # 限制内容长度
                    full_text = "\n\n".join(text_content)
                    if len(full_text) > 100000:
                        full_text = full_text[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    result = f"[PDF文件内容]\n{full_text}"
                    print(f"成功提取PDF文本，总长度: {len(result)}")
                    return result, "text/plain"
                    
                except ImportError:
                    print("未找到PyMuPDF，尝试使用其他PDF处理方法")
                    
                # 尝试方法2: 使用PyPDF2
                try:
                    from PyPDF2 import PdfReader
                    print("使用PyPDF2处理PDF")
                    
                    reader = PdfReader(file_path)
                    text_content = []
                    
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"=== 第 {i + 1} 页 ===\n{text}")
                            print(f"页面 {i + 1} 文本长度: {len(text)}")
                    
                    if not text_content:
                        print("PyPDF2未能提取到文本")
                        return "PDF文件未包含可提取的文本内容，或文本无法被正确提取", "text/plain"
                    
                    # 限制内容长度
                    full_text = "\n\n".join(text_content)
                    if len(full_text) > 100000:  # 约10万字符
                        full_text = full_text[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    result = f"[PDF文件内容]\n{full_text}"
                    print(f"成功提取PDF文本，总长度: {len(result)}")
                    return result, "text/plain"
                    
                except ImportError:
                    print("未找到PyPDF2，尝试使用pdfplumber")
                    
                # 尝试方法3: 使用pdfplumber
                try:
                    import pdfplumber
                    print("使用pdfplumber处理PDF")
                    
                    with pdfplumber.open(file_path) as pdf:
                        text_content = []
                        for i, page in enumerate(pdf.pages):
                            text = page.extract_text()
                            if text and text.strip():
                                text_content.append(f"=== 第 {i + 1} 页 ===\n{text}")
                                print(f"页面 {i + 1} 文本长度: {len(text)}")
                        
                        if not text_content:
                            print("pdfplumber未能提取到文本")
                            return "PDF文件未包含可提取的文本内容，或文本无法被正确提取", "text/plain"
                    
                    # 限制内容长度
                    full_text = "\n\n".join(text_content)
                    if len(full_text) > 100000:  # 约10万字符
                        full_text = full_text[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    result = f"[PDF文件内容]\n{full_text}"
                    print(f"成功提取PDF文本，总长度: {len(result)}")
                    return result, "text/plain"
                    
                except ImportError:
                    print("所有PDF处理库尝试失败")
                    return "未安装PDF处理所需的库，无法提取PDF内容", "text/plain"
                    
            except Exception as e:
                print(f"PDF处理全部方法失败: {str(e)}")
                traceback.print_exc()
                return f"无法处理PDF文件: {str(e)}", "text/plain"
                
        # Word文档处理
        elif ext in ['.docx', '.doc']:
            try:
                # 尝试使用python-docx
                try:
                    import docx
                    print("使用python-docx处理Word文档")
                    
                    # 打开Word文档
                    doc = docx.Document(file_path)
                    
                    # 提取所有段落文本
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    print(f"提取到 {len(paragraphs)} 个段落")
                    
                    # 提取表格内容
                    for i, table in enumerate(doc.tables):
                        table_text = []
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                table_text.append(" | ".join(row_text))
                        if table_text:
                            table_content = f"\n表格 {i+1}:\n" + "\n".join(table_text)
                            paragraphs.append(table_content)
                    
                    if not paragraphs:
                        print("python-docx未能提取到文本内容")
                        return "Word文档未包含可提取的文本内容", "text/plain"
                    
                    # 限制内容长度
                    full_text = "\n\n".join(paragraphs)
                    if len(full_text) > 100000:  # 约10万字符
                        full_text = full_text[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    result = f"[Word文档内容]\n{full_text}"
                    print(f"成功提取Word文本，总长度: {len(result)}")
                    return result, "text/plain"
                    
                except ImportError:
                    print("未找到python-docx，尝试使用textract")
                    
                # 尝试使用textract
                try:
                    import textract
                    print("使用textract处理Word文档")
                    
                    text = textract.process(file_path).decode('utf-8', errors='replace')
                    if not text.strip():
                        print("textract未能提取到文本内容")
                        return "Word文档未包含可提取的文本内容", "text/plain"
                    
                    # 限制内容长度
                    if len(text) > 100000:  # 约10万字符
                        text = text[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    result = f"[Word文档内容]\n{text}"
                    print(f"成功提取Word文本，总长度: {len(result)}")
                    return result, "text/plain"
                    
                except ImportError:
                    print("所有Word处理库尝试失败")
                    return "未安装Word处理所需的库，无法提取Word内容", "text/plain"
                    
            except Exception as e:
                print(f"Word文档处理失败: {str(e)}")
                traceback.print_exc()
                return f"无法处理Word文档: {str(e)}", "text/plain"
                
        # Excel文件处理
        elif ext in ['.xls', '.xlsx']:
            try:
                # 导入pandas和openpyxl
                try:
                    import pandas as pd
                    print("使用pandas处理Excel文件")
                    
                    # 读取所有sheet
                    excel_file = pd.ExcelFile(file_path)
                    sheets = excel_file.sheet_names
                    print(f"Excel文件包含 {len(sheets)} 个工作表")
                    
                    all_sheets_data = []
                    
                    for sheet_name in sheets:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)
                        
                        if df.empty:
                            all_sheets_data.append(f"工作表 '{sheet_name}' 为空")
                            continue
                        
                        # 限制行数，避免过大
                        if len(df) > 100:
                            df = df.head(100)
                            truncated = True
                        else:
                            truncated = False
                        
                        # 将DataFrame转换为字符串格式（表格格式）
                        buffer = io.StringIO()
                        df.to_csv(buffer, sep=',', index=False)
                        csv_text = buffer.getvalue()
                        
                        sheet_text = f"=== 工作表: {sheet_name} ===\n\n" + csv_text
                        
                        if truncated:
                            sheet_text += "\n[数据已截断，仅显示前100行]"
                        
                        all_sheets_data.append(sheet_text)
                    
                    if not all_sheets_data:
                        print("未能从Excel提取到数据")
                        return "Excel文件未包含可提取的数据", "text/plain"
                    
                    result = "[Excel表格数据]\n\n" + "\n\n".join(all_sheets_data)
                    print(f"成功提取Excel数据，总长度: {len(result)}")
                    return result, "text/csv"
                    
                except ImportError as e:
                    print(f"pandas导入失败: {str(e)}")
                    return "未安装Excel处理所需的库 (pandas, openpyxl)", "text/plain"
                    
            except Exception as e:
                print(f"Excel处理失败: {str(e)}")
                traceback.print_exc()
                return f"无法处理Excel文件: {str(e)}", "text/plain"
                
        # CSV文件处理
        elif ext == '.csv':
            try:
                # 检测编码
                encoding = detect_encoding(file_path)
                print(f"检测到CSV文件编码: {encoding}")
                
                try:
                    # 尝试使用pandas
                    import pandas as pd
                    print("使用pandas处理CSV文件")
                    
                    # 尝试不同的分隔符
                    for sep in [',', ';', '\t', '|']:
                        try:
                            df = pd.read_csv(
                                file_path, 
                                encoding=encoding, 
                                sep=sep, 
                                on_bad_lines='skip', 
                                low_memory=False
                            )
                            
                            if not df.empty and df.shape[1] > 1:  # 确保至少有1列数据且分隔正确
                                print(f"使用分隔符 '{sep}' 成功读取CSV")
                                break
                        except Exception as e:
                            print(f"使用分隔符 '{sep}' 读取失败: {str(e)}")
                    
                    if df.empty:
                        print("CSV文件为空或读取失败")
                        return "CSV文件未包含可提取的数据", "text/plain"
                    
                    # 输出表格基本信息
                    print(f"CSV表格尺寸: {df.shape[0]} 行 x {df.shape[1]} 列")
                    
                    # 限制行数，避免过大
                    if len(df) > 500:
                        df = df.head(500)
                        truncated = True
                    else:
                        truncated = False
                    
                    # 转换为字符串格式
                    buffer = io.StringIO()
                    df.to_csv(buffer, sep=',', index=False)
                    csv_text = buffer.getvalue()
                    
                    result = "[CSV表格数据]\n\n" + csv_text
                    
                    if truncated:
                        result += "\n\n[数据已截断，仅显示前500行]"
                    
                    print(f"成功提取CSV数据，总长度: {len(result)}")
                    return result, "text/csv"
                    
                except ImportError:
                    # 如果没有pandas，使用csv模块
                    print("使用csv模块处理CSV文件")
                    import csv
                    
                    rows = []
                    with open(file_path, 'r', encoding=encoding, newline='', errors='replace') as f:
                        # 尝试不同的分隔符
                        for delimiter in [',', ';', '\t', '|']:
                            f.seek(0)  # 重置文件指针
                            try:
                                sample = f.read(1024)
                                f.seek(0)  # 重置文件指针
                                dialect = csv.Sniffer().sniff(sample, delimiters=delimiter)
                                reader = csv.reader(f, dialect)
                                rows = list(reader)
                                if rows and max(len(row) for row in rows) > 1:  # 确保有效的CSV
                                    print(f"使用csv模块和分隔符 '{dialect.delimiter}' 成功读取")
                                    break
                            except Exception as e:
                                print(f"CSV分隔符 '{delimiter}' 检测失败: {str(e)}")
                    
                    if not rows:
                        print("CSV读取失败")
                        return "CSV文件未包含可提取的数据或格式无效", "text/plain"
                    
                    # 限制行数
                    if len(rows) > 500:
                        rows = rows[:500]
                        truncated = True
                    else:
                        truncated = False
                    
                    # 构建文本表示
                    result = []
                    for row in rows:
                        result.append(",".join([f'"{cell}"' if ',' in cell else cell for cell in row]))
                    
                    final_text = "[CSV表格数据]\n\n" + '\n'.join(result)
                    
                    if truncated:
                        final_text += "\n\n[数据已截断，仅显示前500行]"
                    
                    print(f"成功提取CSV数据，总长度: {len(final_text)}")
                    return final_text, "text/csv"
                    
            except Exception as e:
                print(f"CSV处理失败: {str(e)}")
                traceback.print_exc()
                return f"无法处理CSV文件: {str(e)}", "text/plain"
                
        # 文本文件处理
        elif ext in ['.txt', '.md', '.json', '.xml', '.html', '.log', '.py', '.js', '.css']:
            try:
                # 检测编码
                encoding = detect_encoding(file_path)
                print(f"检测到文本文件编码: {encoding}")
                
                # 读取文本文件
                with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                    content = f.read()
                
                if not content.strip():
                    print("文本文件为空")
                    return "文本文件未包含内容", "text/plain"
                
                # 输出文本内容前几行作为调试信息
                lines = content.split('\n')
                preview = '\n'.join(lines[:5]) + ('...' if len(lines) > 5 else '')
                print(f"文本文件前几行: {preview}")
                
                # 限制内容长度
                if len(content) > 100000:  # 约10万字符
                    content = content[:100000] + "\n\n[内容已截断，仅显示前部分]"
                
                file_type = ext[1:].upper()  # 去掉点，转为大写
                result = f"[{file_type}文件内容]\n{content}"
                # print(f"成功提取文本文件内容，总长度: {len(result)}")
                return result, "text/plain"
                
            except Exception as e:
                print(f"文本处理失败: {str(e)}")
                traceback.print_exc()
                
                # 如果用检测到的编码失败，尝试二进制读取
                try:
                    print("尝试二进制读取文本文件")
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='replace')
                    
                    if not content.strip():
                        print("二进制读取结果为空")
                        return "文本文件未包含内容", "text/plain"
                    
                    # 限制内容长度
                    if len(content) > 100000:  # 约10万字符
                        content = content[:100000] + "\n\n[内容已截断，仅显示前部分]"
                    
                    file_type = ext[1:].upper()  # 去掉点，转为大写
                    result = f"[{file_type}文件内容 - 二进制读取]\n{content}"
                    print(f"成功通过二进制读取文本文件内容，总长度: {len(result)}")
                    return result, "text/plain"
                except Exception as inner_e:
                    print(f"二次尝试文本处理失败: {str(inner_e)}")
                    return f"无法读取文本文件: {str(e)}", "text/plain"
                    
        # 其他文件类型
        else:
            print(f"不支持的文件类型: {ext}")
            return f"暂不支持处理文件类型 {ext}，请上传 PDF、Word、Excel、CSV、文本或图像文件。", "text/plain"
    
    except Exception as e:
        print(f"处理文件 {filename} 时出错: {str(e)}")
        traceback.print_exc()
        return f"处理文件时发生错误: {str(e)}", "text/plain"